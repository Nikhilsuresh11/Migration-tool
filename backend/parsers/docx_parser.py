
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from utils.helpers import clean_text


def parse_docx(filepath: str) -> dict:

    try:
        doc = Document(filepath)
    except PackageNotFoundError:
        raise ValueError("The file is not a valid .docx document or is corrupted.")
    except Exception as e:
        raise ValueError(f"Failed to open .docx file: {str(e)}")

    headings = extract_headings(doc)
    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)
    full_text = extract_full_text(doc)
    images_count = count_images(doc)

    pages = [{
        "page_number": 1,
        "text": full_text,
        "word_count": len(full_text.split()) if full_text else 0,
    }]

    return {
        "file_type": "docx",
        "full_text": full_text,
        "headings": headings,
        "paragraphs": paragraphs,
        "pages": pages,
        "tables": tables,
        "images_count": images_count,
        "metadata": extract_metadata(doc),
    }


def extract_headings(doc: Document) -> list:

    headings = []
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()
        if style_name.startswith("heading"):
            level = _parse_heading_level(style_name)
            text = paragraph.text.strip()
            if text:
                headings.append({"text": text, "level": level})
    return headings


def _parse_heading_level(style_name: str) -> int:
    parts = style_name.split()
    if len(parts) >= 2 and parts[-1].isdigit():
        return int(parts[-1])
    return 1


def extract_paragraphs(doc: Document) -> list:

    paragraphs = []
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()
        text = paragraph.text.strip()
        if text and not style_name.startswith("heading"):
            paragraphs.append(clean_text(text))
    return paragraphs


def extract_tables(doc: Document) -> list:

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def extract_full_text(doc: Document) -> str:

    text_parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    return clean_text("\n\n".join(text_parts))


def count_images(doc: Document) -> int:
    count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            count += 1
    return count


def extract_metadata(doc: Document) -> dict:

    props = doc.core_properties
    return {
        "author": props.author or "",
        "title": props.title or "",
        "subject": props.subject or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
    }
