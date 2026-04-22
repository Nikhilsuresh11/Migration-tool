"""
Table quality analysis module.

Analyzes table structure and quality for migration readiness,
including header detection, merged cell detection (DOCX),
complex table flagging, and improved PDF table extraction
using PyMuPDF's find_tables() method.
"""

from docx import Document
import fitz


def analyze_tables(filepath, file_type, parsed_tables=None, warnings=None):
    """
    Analyze table quality and structure in a document.

    Args:
        filepath: Path to the document file.
        file_type: 'docx' or 'pdf'.
        parsed_tables: Tables already extracted by the parser (for DOCX).
        warnings: Shared list to append warning messages to.

    Returns:
        Dictionary with table_count, tables_with_headers,
        merged_cells_detected, and complex_tables counts.
    """
    if warnings is None:
        warnings = []

    try:
        if file_type == "docx":
            return _analyze_docx_tables(filepath)
        elif file_type == "pdf":
            return _analyze_pdf_tables(filepath)
        return get_default()
    except Exception as e:
        warnings.append(f"Table analysis failed: {str(e)}")
        return get_default()


def _analyze_docx_tables(filepath):
    """Analyze tables in a DOCX file for quality metrics."""
    doc = Document(filepath)

    table_count = len(doc.tables)
    tables_with_headers = 0
    merged_cells = 0
    complex_tables = 0

    for table in doc.tables:
        rows = table.rows
        cols_count = len(rows[0].cells) if rows else 0
        row_count = len(rows)

        # Check for headers (first row looks like a header)
        if rows and _is_header_row_docx(rows[0]):
            tables_with_headers += 1

        # Detect merged cells
        merged_cells += _count_merged_cells_docx(table)

        # Flag complex tables (> 6 columns or > 20 rows)
        if cols_count > 6 or row_count > 20:
            complex_tables += 1

    return {
        "table_count": table_count,
        "tables_with_headers": tables_with_headers,
        "merged_cells_detected": merged_cells,
        "complex_tables": complex_tables,
    }


def _is_header_row_docx(row):
    """
    Determine if a DOCX table row looks like a header row.

    A header row typically has non-empty cells with distinct text
    and is often formatted differently (bold text).
    """
    cells = row.cells
    if not cells:
        return False

    # Check if all cells have text
    cell_texts = [cell.text.strip() for cell in cells]
    non_empty = [t for t in cell_texts if t]

    if not non_empty:
        return False

    # Check if first row cells contain bold text (common header formatting)
    has_bold = False
    try:
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        has_bold = True
                        break
    except Exception:
        pass

    # Consider it a header if most cells have text and/or bold formatting
    if len(non_empty) >= len(cells) * 0.5:
        return True

    return has_bold


def _count_merged_cells_docx(table):
    """
    Count merged cells in a DOCX table.

    Merged cells in python-docx share the same underlying _tc element.
    We detect this by tracking unique _tc element IDs per row.
    """
    merged = 0
    for row in table.rows:
        seen_tcs = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                merged += 1
            else:
                seen_tcs.add(tc_id)
    return merged


def _analyze_pdf_tables(filepath):
    """
    Analyze tables in a PDF file using PyMuPDF's find_tables().

    This provides much better table detection than the previous
    approach which always returned 0 for PDFs.
    """
    doc = fitz.open(filepath)

    table_count = 0
    tables_with_headers = 0
    complex_tables = 0

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)

        try:
            tabs = page.find_tables()
        except AttributeError:
            # find_tables() not available in this PyMuPDF version
            doc.close()
            return get_default()

        for table in tabs.tables:
            table_count += 1

            row_count = table.row_count
            col_count = table.col_count

            # Check for header
            if _is_header_row_pdf(table):
                tables_with_headers += 1

            # Flag complex tables
            if col_count > 6 or row_count > 20:
                complex_tables += 1

    doc.close()

    return {
        "table_count": table_count,
        "tables_with_headers": tables_with_headers,
        "merged_cells_detected": 0,  # Merged cell detection not reliable in PDF
        "complex_tables": complex_tables,
    }


def _is_header_row_pdf(table):
    """
    Check if a PDF table has a detectable header row.

    Uses PyMuPDF Table's header property if available,
    otherwise checks if the first row has content.
    """
    try:
        # PyMuPDF Table.header contains header info
        header = table.header
        if header and hasattr(header, "names"):
            names = header.names
            non_empty = [n for n in names if n and str(n).strip()]
            return len(non_empty) > 0
    except Exception:
        pass

    # Fallback: check if first row of extracted data has text
    try:
        data = table.extract()
        if data and data[0]:
            first_row = data[0]
            non_empty = [c for c in first_row if c and str(c).strip()]
            return len(non_empty) >= len(first_row) * 0.5
    except Exception:
        pass

    return False


def get_default():
    """Return the default empty structure for table analysis."""
    return {
        "table_count": 0,
        "tables_with_headers": 0,
        "merged_cells_detected": 0,
        "complex_tables": 0,
    }
